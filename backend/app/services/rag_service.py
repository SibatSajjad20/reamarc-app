import logging
import math
import re
import asyncio
from typing import List, Dict, Any, Optional
import fitz  # PyMuPDF
import docx  # python-docx
from bs4 import BeautifulSoup
from app.config import settings

logger = logging.getLogger(__name__)

# Preferred local open-source embedding model (384 dimensions, zero rate limit, $0 cost)
LOCAL_EMBEDDING_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"
_fastembed_model = None


def extract_text_from_pdf(contents: bytes) -> tuple[str, int]:
    """Extract text and page count from PDF bytes using PyMuPDF."""
    doc = fitz.open(stream=contents, filetype="pdf")
    extracted = ""
    for page in doc:
        extracted += page.get_text() + "\n"
    page_count = len(doc)
    doc.close()
    return extracted.strip(), page_count


def extract_text_from_docx(contents: bytes) -> tuple[str, int]:
    """Extract text from Word .docx bytes using python-docx."""
    import io
    doc = docx.Document(io.BytesIO(contents))
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
    
    extracted = "\n".join(full_text)
    section_count = max(1, len(full_text))
    return extracted.strip(), section_count


def extract_text_from_txt(contents: bytes) -> tuple[str, int]:
    """Extract text from raw text / markdown / csv bytes."""
    try:
        text = contents.decode("utf-8")
    except UnicodeDecodeError:
        text = contents.decode("latin-1", errors="ignore")
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    return "\n".join(lines), len(lines)


def extract_text_from_html(html_content: str) -> str:
    """Clean and extract core content text from HTML webpage string."""
    soup = BeautifulSoup(html_content, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
        tag.decompose()
    extracted_text = soup.get_text(separator="\n", strip=True)
    lines = [line.strip() for line in extracted_text.splitlines() if line.strip()]
    return "\n".join(lines)


def chunk_text(text: str, max_chunk_size: int = 600, overlap: int = 80) -> List[str]:
    """
    Sentence-aware recursive text chunker.
    Splits text into logical paragraphs/sentences aiming for max_chunk_size with overlap.
    """
    clean_text = text.strip()
    if not clean_text:
        return []

    if len(clean_text) <= max_chunk_size:
        return [clean_text]

    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', clean_text) if p.strip()]
    chunks: List[str] = []
    current_chunk = ""

    for p in paragraphs:
        if len(p) > max_chunk_size:
            sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', p) if s.strip()]
            for s in sentences:
                if len(current_chunk) + len(s) + 1 <= max_chunk_size:
                    current_chunk = f"{current_chunk} {s}".strip()
                else:
                    if current_chunk:
                        chunks.append(current_chunk)
                    overlap_prefix = current_chunk[-overlap:] if len(current_chunk) > overlap else ""
                    current_chunk = f"{overlap_prefix} {s}".strip()
        else:
            if len(current_chunk) + len(p) + 2 <= max_chunk_size:
                current_chunk = f"{current_chunk}\n\n{p}".strip()
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                overlap_prefix = current_chunk[-overlap:] if len(current_chunk) > overlap else ""
                current_chunk = f"{overlap_prefix}\n\n{p}".strip()

    if current_chunk:
        chunks.append(current_chunk)

    return chunks


def _get_local_embedding_model():
    global _fastembed_model
    if _fastembed_model is not None:
        return _fastembed_model
    try:
        from fastembed import TextEmbedding
        _fastembed_model = TextEmbedding(model_name=LOCAL_EMBEDDING_MODEL_NAME)
        logger.info(f"Initialized local embedding model: {LOCAL_EMBEDDING_MODEL_NAME}")
        return _fastembed_model
    except Exception as err:
        logger.warning(f"Could not load local embedding model '{LOCAL_EMBEDDING_MODEL_NAME}': {err}")
        return None


def _fallback_hash_vector(text: str, dimensions: int = 384) -> List[float]:
    """Deterministic term-frequency hashing vectorizer fallback when offline."""
    words = re.findall(r'\w+', text.lower())
    vec = [0.0] * dimensions
    if not words:
        return vec
    for word in words:
        idx = abs(hash(word)) % dimensions
        vec[idx] += 1.0
    norm = math.sqrt(sum(val * val for val in vec))
    if norm > 0:
        vec = [val / norm for val in vec]
    return vec


async def generate_embedding(text: str) -> List[float]:
    """Generate 384D float vector embedding for input text locally using sentence-transformers/all-MiniLM-L6-v2."""
    if not text or not text.strip():
        return _fallback_hash_vector("empty", 384)

    try:
        model = _get_local_embedding_model()
        if model:
            loop = asyncio.get_running_loop()

            def _embed_sync(t_str: str) -> List[float]:
                embeddings = list(model.embed([t_str[:2000]]))
                if embeddings and len(embeddings) > 0:
                    return embeddings[0].tolist()
                return []

            vec = await loop.run_in_executor(None, _embed_sync, text)
            if vec and len(vec) > 0:
                return vec
    except Exception as e:
        logger.warning(f"Local embedding generation failed: {e}")

    return _fallback_hash_vector(text, 384)


def cosine_similarity(vec_a: List[float], vec_b: List[float]) -> float:
    """Calculate cosine similarity between two float vectors."""
    if not vec_a or not vec_b or len(vec_a) != len(vec_b):
        return 0.0
    dot = sum(a * b for a, b in zip(vec_a, vec_b))
    norm_a = math.sqrt(sum(a * a for a in vec_a))
    norm_b = math.sqrt(sum(b * b for b in vec_b))
    if norm_a == 0.0 or norm_b == 0.0:
        return 0.0
    return dot / (norm_a * norm_b)


async def index_document_chunks(
    db,
    user_id: str,
    workspace_id: str,
    source_id: str,
    source_name: str,
    source_type: str,
    full_text: str
) -> int:
    """
    Chunk document text, compute embeddings asynchronously, and insert chunks into MongoDB 'knowledge_chunks'.
    Returns total chunk count indexed.
    """
    chunks = chunk_text(full_text, max_chunk_size=600, overlap=80)
    if not chunks:
        return 0

    chunk_docs = []
    for idx, chunk_str in enumerate(chunks):
        embedding = await generate_embedding(chunk_str)
        chunk_docs.append({
            "id": f"chunk-{source_id}-{idx}",
            "source_id": source_id,
            "workspaceId": workspace_id,
            "user_id": user_id,
            "chunk_index": idx,
            "text": chunk_str,
            "embedding": embedding,
            "source_name": source_name,
            "source_type": source_type,
        })

    if chunk_docs:
        await db.knowledge_chunks.insert_many(chunk_docs)

    return len(chunk_docs)


async def retrieve_relevant_chunks(
    db,
    user_id: str,
    workspace_id: str,
    query_text: str,
    top_k: int = 6
) -> List[Dict[str, Any]]:
    """
    Perform semantic vector similarity search against workspace chunks stored in MongoDB.
    Returns top_k most relevant chunks.
    """
    if db is None or not query_text.strip():
        return []

    try:
        # Generate query vector embedding
        query_embedding = await generate_embedding(query_text)

        # Retrieve candidate chunks for this user & workspace
        cursor = db.knowledge_chunks.find(
            {"workspaceId": workspace_id, "user_id": user_id},
            {"_id": 0, "text": 1, "embedding": 1, "source_name": 1, "source_type": 1, "chunk_index": 1}
        )
        candidate_chunks = await cursor.to_list(length=500)

        if not candidate_chunks:
            return []

        scored_chunks = []
        for chunk in candidate_chunks:
            embedding = chunk.get("embedding")
            if embedding:
                score = cosine_similarity(query_embedding, embedding)
                scored_chunks.append({
                    "score": score,
                    "text": chunk.get("text", ""),
                    "source_name": chunk.get("source_name", "Document"),
                    "source_type": chunk.get("source_type", "pdf"),
                })

        # Sort descending by similarity score
        scored_chunks.sort(key=lambda x: x["score"], reverse=True)
        return scored_chunks[:top_k]
    except Exception as err:
        logger.warning(f"Error during RAG semantic search retrieval: {err}")
        return []
